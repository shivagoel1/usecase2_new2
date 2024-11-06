import requests
import streamlit as st
import os
import warnings
from crewai import Agent, Task, Crew
import io
from docx import Document
from docx.shared import Pt, RGBColor

# Suppress warnings
warnings.filterwarnings('ignore')

# Streamlit UI
st.title("Research Article Generator")

# Multiple file uploader
uploaded_files = st.file_uploader("Upload one or more transcript files", type="txt", accept_multiple_files=True)

# Display the file names concisely
if uploaded_files:
    file_names = [file.name for file in uploaded_files]
    st.write("Uploaded Files:", ", ".join(file_names))

# API Key input
openai_api_key = st.text_input("Enter your OpenAI API Key", type="password")

# Button to start processing
if st.button("Generate Research Article"):
    if not uploaded_files:
        st.error("Please upload at least one transcript file.")
    elif not openai_api_key:
        st.error("Please enter your OpenAI API Key.")
    else:
        # Set up environment variables
        os.environ["OPENAI_API_KEY"] = openai_api_key
        os.environ["OPENAI_MODEL_NAME"] = 'gpt-4o'

        # Concatenate all file contents into a single string
        transcripts = ""
        for uploaded_file in uploaded_files:
            file_content = uploaded_file.read().decode("utf-8", errors="ignore")  # Ensure UTF-8 decoding
            transcripts += f"\n{file_content}"

        # Test API connection
        try:
            response = requests.get("https://api.openai.com/v1/models",
                headers={"Authorization": f"Bearer {openai_api_key}"})
            response.raise_for_status()  # Raise an error for bad responses
            st.success("API connection successful!")

            # Define agents with original prompts
            planner = Agent(
                role="Content Planner",
                goal="Plan engaging and factually accurate content on the given topic",
                backstory=(
                    "You're responsible for analyzing the transcripts to extract key themes, challenges, "
                    "and opportunities discussed by industry leaders. Categorize the insights into major "
                    "sections, such as Industry Trends, Technological Impacts, Regulatory Considerations, and Future Outlook. "
                    "Use participant quotes strategically to add credibility and depth, ensuring you include specific examples "
                    "from relevant companies where applicable. "
                    "Ensure the report reads naturally and has the polished "
                    "feel of a human-written document, with varied sentence structures, a professional tone, and engaging, nuanced language."
                ),
                allow_delegation=False,
                verbose=True
            )

            writer = Agent(
                role="Content Writer",
                goal="Write insightful and factually accurate research report about the given topic",
                backstory=(
                    "Your task is to write a comprehensive and engaging research article based on the content "
                    "plan provided by the Content Planner. Integrate specific quotes from participants to support "
                    "key arguments and provide a balanced view of the opportunities and challenges discussed. "
                    "Use evidence-based analysis and maintain a formal yet engaging tone. Structure the content "
                    "thematically, addressing each major point with supporting data, expert opinions, and specific "
                    "examples. Highlight knowledge gaps and propose strategies for addressing them, ensuring the content "
                    "is actionable. Write in a way that feels human and natural, as though crafted by a seasoned technical "
                    "writer. Avoid robotic language and ensure the narrative is engaging, relatable, and enriched with "
                    "cross-references that connect different sections of the report for a cohesive flow. "
                    "End the article with a final 'Conclusion' section, which summarizes key insights without adding further suggestions or recommendations."
                ),
                allow_delegation=False,
                verbose=True
            )

            editor = Agent(
                role="Editor",
                goal="Edit a given research article to align with the writing style of the organization",
                backstory=(
                    "Your role is to refine the research article drafted by the Content Writer. Ensure the content "
                    "follows journalistic best practices, maintains a formal and professional tone, and is well-structured. "
                    "Check for balanced viewpoints and make sure that participant quotes are used effectively. Avoid "
                    "controversial statements unless necessary, and ensure the report addresses both benefits and risks. "
                    "Focus on coherence, readability, and the logical flow of ideas. Make sure there is no content or "
                    "additional sections following the Conclusion. The Conclusion should be the final part of the report, "
                    "summarizing key insights without adding any further recommendations or suggestions."
                ),
                allow_delegation=False,
                verbose=True
            )

            # Define tasks with original descriptions
            plan = Task(
                description=(
                    "Analyze the transcripts to extract major themes and plan the content structure. Identify key challenges, "
                    "opportunities, and knowledge gaps, and suggest where to include participant quotes. Recommend specific case studies, "
                    "examples, or statistics that would enrich the report."
                ),
                agent=planner,
                expected_output=(
                    "A detailed content outline with categorized themes, key insights, strategic use of quotes, and recommendations "
                    "for case studies"
                )
            )

            write = Task(
                description=(
                    "Write a research article based on the content plan, integrating participant quotes, evidence-based analysis, specific examples, "
                    "and a balanced discussion of opportunities and risks. Ensure the content is engaging, relatable, and structured to connect different themes. "
                    "End the article with a final 'Conclusion' section, which summarizes the report without adding further suggestions or recommendations."
                ),
                agent=writer,
                expected_output=(
                    "A well-written and insightful research article that follows the content plan and addresses all major themes comprehensively, "
                    "with humanized language and cross-references. Ensure there is no content after the Conclusion section."
                )
            )

            edit = Task(
                description=(
                    "Review and edit the research article to ensure coherence, proper use of quotes, balanced viewpoints, and adherence to journalistic standards. "
                    "Make sure that cross-references are present and that the article ends with a Conclusion section only, with no additional recommendations or suggestions afterward."
                ),
                agent=editor,
                expected_output="A polished and professional research article that is ready for publication."
            )

            # Create crew and add tasks
            crew = Crew(
                agents=[planner, writer, editor],
                tasks=[plan, write, edit],
                verbose=True
            )

            # Process the transcript
            with st.spinner("Generating research article... This may take a few minutes."):
                result = crew.kickoff()  # Execute the process

            # Attempt to extract and display the content from the result
            text_content = result.raw  # Ensure you access the correct attribute for the text
            text_content = text_content.encode('utf-8', errors='ignore').decode('utf-8')  # Clean unsupported characters
            st.success("Research article generated successfully!")
            st.markdown(text_content)  # Display the content as Markdown

            # Generate Word document with specified formatting
            doc = Document()
            
            # Set document margins to 1 inch
            doc_sections = doc.sections
            for section in doc_sections:
                section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Pt(72)  # 1 inch margin
            
            # Add content with navy blue color for subheadings
            doc.add_paragraph("Industry Insights Report", style='Heading 1')
            
            # Assuming specific keywords can identify subheadings
            subheading_keywords = ["Industry Trends", "Technological Impacts", "Regulatory Considerations", "Future Outlook", "Conclusion"]

            for line in text_content.split('\n'):
                clean_line = line.strip('*')  # Remove asterisks from each line
                if any(keyword in clean_line for keyword in subheading_keywords):  # Check if the line is a subheading
                    p = doc.add_paragraph(clean_line)
                    p.style.font.name = 'Times New Roman'
                    p.style.font.size = Pt(11)
                    p.runs[0].font.color.rgb = RGBColor(0, 0, 128)  # Navy blue color for subheadings
                else:
                    p = doc.add_paragraph(clean_line)
                    p.style.font.name = 'Times New Roman'
                    p.style.font.size = Pt(11)
                p.paragraph_format.alignment = 0  # Left align
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1  # Single line spacing

            # Save the document to a buffer
            word_buffer = io.BytesIO()
            doc.save(word_buffer)
            word_buffer.seek(0)

            # Download Word document
            st.download_button(
                label="Download Word Document",
                data=word_buffer.getvalue(),
                file_name="research_article.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        except requests.exceptions.RequestException as e:
            st.error(f"API Error: {str(e)}")
            if hasattr(e, 'response'):
                st.error(f"Response Status Code: {e.response.status_code}")
                st.error(f"Response Content: {e.response.text}")
        except Exception as e:
            st.error(f"An error occurred: {str(e)}")

# Footer
st.markdown("---")
st.markdown("Tapestry Networks")